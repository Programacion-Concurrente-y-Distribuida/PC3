#!/usr/bin/env python3
"""Convierte un markdown simple (titulos, parrafos, tablas, codigo, negritas) a .docx.

Uso: python3 md_to_docx.py entrada.md salida.docx
"""

import re
import sys

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_runs(par, text):
    # Reparte el texto en runs normales/negrita segun ** y elimina backticks.
    for i, chunk in enumerate(text.split("**")):
        chunk = chunk.replace("`", "")
        if not chunk:
            continue
        run = par.add_run(chunk)
        run.bold = i % 2 == 1


def add_table(doc, lines):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        rows.append(cells)
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            par = cell.paragraphs[0]
            if i == 0:
                par.add_run(val.replace("**", "").replace("`", "")).bold = True
            else:
                add_runs(par, val)
                if j > 0:
                    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()


def convert(md_path, docx_path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    lines = open(md_path, encoding="utf-8").read().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            par = doc.add_paragraph()
            run = par.add_run("\n".join(block))
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            i += 1
            continue

        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, block)
            continue

        if m := re.match(r"^(#{1,4})\s+(.*)", line):
            doc.add_heading(m.group(2).replace("`", ""), level=len(m.group(1)))
            i += 1
            continue

        if m := re.match(r"^[-*]\s+(.*)", line):
            par = doc.add_paragraph(style="List Bullet")
            add_runs(par, m.group(1))
            i += 1
            continue

        if line.strip():
            # Une lineas consecutivas del mismo parrafo.
            block = [line]
            while (
                i + 1 < len(lines)
                and lines[i + 1].strip()
                and not re.match(r"^(#|\||```|[-*]\s)", lines[i + 1])
            ):
                i += 1
                block.append(lines[i])
            add_runs(doc.add_paragraph(), " ".join(block))
        i += 1

    doc.save(docx_path)
    print(f"OK: {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])

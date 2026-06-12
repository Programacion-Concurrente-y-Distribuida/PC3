#!/usr/bin/env python3
"""Genera diseno_concurrencia.docx (version condensada para el informe)."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h(text, level=1):
    doc.add_heading(text, level=level)


def p(text, bold_prefix=None):
    par = doc.add_paragraph()
    if bold_prefix:
        par.add_run(bold_prefix).bold = True
    par.add_run(text)
    return par


def bullet(text, bold_prefix=None):
    par = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        par.add_run(bold_prefix).bold = True
    par.add_run(text)


def mono(text):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def table(headers, rows, bold_first_col=False, bold_row=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, htxt in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        cell.paragraphs[0].add_run(htxt).bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            if (bold_first_col and j == 0) or (bold_row is not None and i - 1 == bold_row):
                run.bold = True
            if j > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()


h("Modelo de concurrencia y arquitectura del sistema", 1)

h("1. Detección del cuello de botella", 2)
p(
    "La concurrencia no se aplicó a ciegas: un modo de profiling (-profile) instrumenta con logs cada "
    "etapa del flujo, midiendo tiempo, memoria y ciclos de GC. Sobre los 3M de registros, los logs "
    "revelaron dos puntos calientes: la carga y procesamiento del CSV (load_csv_pipeline, ~5.9 s, "
    "~58 % del total) y el ajuste del modelo (acumular XᵀX y Xᵀy sobre 1.88M de filas). La "
    "concurrencia se aplicó exactamente ahí; las etapas baratas (split, métricas, predicción) se "
    "mantuvieron secuenciales."
)
mono(
    "[profile] end  load_csv_pipeline     elapsed=  5.887s\n"
    "[profile] end  standardize           elapsed=  2.866s\n"
    "[profile] end  fit_linear_regression elapsed=  0.268s\n"
    "[profile] end  total                 elapsed= 10.142s"
)

h("2. Tipo de concurrencia", 2)
bullet(
    "pipeline productor–consumidor con pools de workers (fan-out/fan-in). Una "
    "goroutine lee el CSV (el disco es secuencial y csv.Reader no es seguro para concurrencia), un "
    "pool de N parsers valida y parsea, y un pool de M encoders codifica las categóricas. Cada pool "
    "consume de un canal y publica en el siguiente.",
    bold_prefix="Ingesta del CSV: ",
)
bullet(
    "paralelismo de datos tipo map-reduce. Las filas se parten en bloques "
    "contiguos; cada goroutine acumula sus matrices parciales XᵀX y Xᵀy en memoria local, una "
    "WaitGroup actúa de barrera y la reducción suma los parciales antes de resolver por Cholesky "
    "(con respaldo SVD). No hay escrituras compartidas durante el cómputo: las condiciones de "
    "carrera se eliminan por diseño, sin mutexes.",
    bold_prefix="Entrenamiento: ",
)

h("3. Canales y buffers", 2)
table(
    ["Canal", "Conecta", "Buffer"],
    [
        ["rawRowsCh", "lectora → parsers", "4 (def.: 2×parsers)"],
        ["parsedCh", "parsers → encoders", "4 (def.: 2×encoders)"],
        ["encodedCh", "encoders → recolector", "3"],
    ],
    bold_first_col=True,
)
bullet(
    "sin buffer, cada etapa se sincronizaría fila a fila, serializando el "
    "pipeline; el buffer permite que todas las etapas trabajen a la vez.",
    bold_prefix="Por qué con buffer: ",
)
bullet(
    "actúan como control de flujo (backpressure): si una etapa se atrasa, el canal "
    "se llena y la anterior se bloquea sola al enviar, acotando la memoria en tránsito sin "
    "semáforos manuales.",
    bold_prefix="Por qué pequeños: ",
)
bullet(
    "cada etapa cierra su canal de salida al terminar (EOF en la lectora; "
    "WaitGroup + close en cada pool) y el cierre se propaga en cascada hasta el recolector.",
    bold_prefix="Señal de fin: ",
)
p(
    "Otras primitivas: contadores atomic.Int64 para filas cargadas/omitidas (sin locks), "
    "context.WithCancel para la parada cooperativa de todo el pipeline (todos los bucles escuchan "
    "ctx.Done()), y copia defensiva de cada registro porque csv.Reader reutiliza su array interno."
)

h("4. Cómo se determinó el número de goroutines", 2)
p(
    "El número de fit-workers se fijó con un barrido experimental: 8 configuraciones (4 a 32 "
    "goroutines) × 20 repeticiones cada una sobre el dataset completo (2,285,426 filas), en una "
    "máquina de 14 núcleos, midiendo tiempo total, tiempo de ajuste y RAM pico."
)
table(
    ["fit-workers", "fit (s, med.)", "total (s, med.)", "desv. est. (s)", "RAM (MB)"],
    [
        ["4", "0.813", "14.17", "2.38", "5,309"],
        ["8", "0.407", "17.50", "2.37", "4,329"],
        ["12", "0.333", "15.18", "2.59", "5,127"],
        ["16", "0.421", "17.31", "39.24", "4,769"],
        ["20", "0.347", "14.88", "14.52", "5,281"],
        ["24", "0.299", "10.89", "1.06", "5,856"],
        ["28", "0.280", "10.11", "0.55", "5,858"],
        ["32", "0.284", "10.40", "0.44", "5,592"],
    ],
    bold_row=6,
)
bullet(
    "el fit cae de 0.81 s (4 workers) a 0.28 s (28): speedup ~2.9×. La curva se "
    "aplana desde ~24 workers al saturarse los núcleos físicos.",
    bold_prefix="Escalamiento: ",
)
bullet(
    "28 goroutines (2× los 14 núcleos): menor mediana total (10.11 s), menor fit y "
    "ejecución estable (desv. 0.55 s). El doble de workers que núcleos compensa bloqueos y "
    "desalojos del planificador.",
    bold_prefix="Óptimo: ",
)
bullet(
    "con 32 workers el tiempo ya no mejora y la RAM pico sube (~5.3 → ~5.9 GB), "
    "porque cada worker preasigna sus parciales; eso fija el techo práctico.",
    bold_prefix="Más no es gratis: ",
)
bullet(
    "las métricas (MAE 0.7755, RMSE 1.0522, R² 0.9004) fueron idénticas en las "
    "160 corridas con cualquier número de workers: evidencia empírica de ausencia de condiciones "
    "de carrera — la concurrencia cambia el tiempo, nunca el resultado.",
    bold_prefix="Corrección: ",
)
p(
    "Para la ingesta existe un benchmark análogo de la grilla parser×encoder workers (mapas de calor "
    "de tiempo y RAM); allí los retornos decrecen antes porque la etapa lectora es I/O secuencial: el "
    "límite no es cuántos parsers haya, sino la velocidad del disco."
)

h("5. Síntesis", 2)
p(
    "La arquitectura sigue un principio único: compartir comunicando (canales) en lugar de comunicar "
    "compartiendo (memoria + locks). Donde el flujo es continuo se usa pipeline con backpressure; "
    "donde el cómputo es una suma particionable, map-reduce con estado local y barrera. El sistema "
    "procesa 3 millones de registros y entrena sobre 1.88 millones de filas en ~10 segundos, de forma "
    "determinista y con memoria acotada."
)

out = "/Users/azariel/Documents/Github/PC3/diseno_concurrencia.docx"
doc.save(out)
print(f"OK: {out}")

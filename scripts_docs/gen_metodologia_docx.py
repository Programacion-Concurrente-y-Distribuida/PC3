#!/usr/bin/env python3
"""Genera metodologia_datos_y_modelo.docx (version condensada para el informe)."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h(text, level=1):
    doc.add_heading(text, level=level)


def p(text, bold_prefix=None):
    par = doc.add_paragraph()
    if bold_prefix:
        par.add_run(bold_prefix).bold = True
    par.add_run(text)
    return par


def bullet(text, bold_prefix=None):
    par = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        par.add_run(bold_prefix).bold = True
    par.add_run(text)


def table(headers, rows, bold_first_col=False):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, htxt in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(htxt)
        run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            if bold_first_col and j == 0:
                run.bold = True
            if j > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()


h("Metodología: datos, modelo y verificación de métricas", 1)

h("1. Variable objetivo y uso de los datos", 2)
p(
    "Arithmetic Mean, la concentración media anual del contaminante por estación "
    "de monitoreo del sistema AQS de la EPA. Predecirla permite estimar la exposición "
    "de una zona sin contar con la serie completa de mediciones.",
    bold_prefix="Variable a predecir: ",
)
p(
    "El insumo es el CSV de la Fase 1 (aqs_final_3M.csv, 3 millones de registros). "
    "El pipeline concurrente valida cada fila (numéricos parseables, sin vacíos, NaN ni "
    "infinitos): 2,285,426 filas válidas cargadas y 714,574 omitidas."
)
bullet(
    "11 — Latitude, Longitude, Year, Observation Count, Observation Percent, "
    "Valid Day Count, Required Day Count, Null Data Count, Num Obs Below MDL, "
    "Primary y Secondary Exceedance Count.",
    bold_prefix="Features numéricas: ",
)
bullet(
    "pollutant, Sample Duration y State Code, con one-hot encoding. "
    "Total tras codificar: 70 features.",
    bold_prefix="Features categóricas: ",
)
bullet(
    "máximos anuales y percentiles, por derivar de la misma distribución que el "
    "target (data leakage).",
    bold_prefix="Excluidas: ",
)
p(
    "temporal, no aleatoria. Años ≤ 2020 entrenan (1,880,537 filas) y años > 2020 "
    "prueban (404,889). Simula el uso real (entrenar con historia, predecir el futuro) y es más "
    "exigente que un split aleatorio. La estandarización se calcula solo con el train y se aplica "
    "al test, evitando fuga de información.",
    bold_prefix="Partición train/test: ",
)

h("2. Modelo elegido y justificación", 2)
p(
    "Regresión lineal múltiple con regularización Ridge (λ = 1, sin penalizar el intercepto), "
    "implementada en Go sobre gonum. Razones:"
)
bullet("los coeficientes indican qué variables pesan más en la predicción.", bold_prefix="Interpretabilidad: ")
bullet(
    "el ajuste por ecuaciones normales se reduce a acumular XᵀX (70×70) y Xᵀy, "
    "paralelizable por bloques de filas entre goroutines.",
    bold_prefix="Escalabilidad: ",
)
bullet(
    "Ridge estabiliza la colinealidad del one-hot; si Cholesky detecta matriz "
    "singular, el sistema cae automáticamente a un solver SVD.",
    bold_prefix="Robustez: ",
)
p(
    "Se evaluó además PCA (90 %, 95 % y 99 % de varianza) como etapa opcional, para verificar "
    "con datos si la reducción de dimensionalidad aporta."
)

h("3. Métricas y protocolo de verificación", 2)
table(
    ["Métrica", "Qué mide", "Por qué"],
    [
        ["MAE", "Error absoluto promedio", "Interpretable en unidades del contaminante; robusto a outliers"],
        ["RMSE", "Raíz del error cuadrático medio", "Penaliza errores grandes"],
        ["R²", "Varianza explicada", "Escala común (0–1) para comparar escenarios"],
    ],
    bold_first_col=True,
)
bullet(
    "cada escenario se ejecutó 3 veces con un script automatizado. La desviación "
    "estándar de las métricas fue 0.0 en todos: el pipeline es determinista — la concurrencia "
    "cambia el tiempo, no el resultado.",
    bold_prefix="Repetición: ",
)
bullet(
    "R² train 0.9207 vs R² test 0.9004; brecha pequeña, buena "
    "generalización a años no vistos.",
    bold_prefix="Control de overfitting: ",
)
bullet(
    "cada variante PCA se compara contra el baseline por deltas de MAE, RMSE y R², "
    "más tiempo total y RAM pico.",
    bold_prefix="Comparación: ",
)

h("4. Resultados (test, promedio de 3 corridas)", 2)
table(
    ["Escenario", "MAE", "RMSE", "R²", "Tiempo (s)", "RAM (MB)", "Componentes"],
    [
        ["sin_pca", "0.7755", "1.0522", "0.9004", "103.7", "4,055", "—"],
        ["pca_90", "1.0174", "1.2987", "0.8483", "88.6", "4,876", "53"],
        ["pca_95", "1.0244", "1.2941", "0.8494", "46.1", "4,914", "57"],
        ["pca_99", "0.8510", "1.1532", "0.8804", "68.0", "4,701", "60"],
    ],
    bold_first_col=True,
)
bullet("el baseline sin PCA gana en las tres métricas; toda variante PCA degrada la calidad (ΔR² de −0.020 a −0.052).")
bullet(
    "con solo 70 features, PCA descarta poca dimensionalidad real y diluye la señal de las variables "
    "dummy categóricas; pca_99 (60/70 componentes) es la menos dañina, confirmando que la reducción no aporta."
)
bullet(
    "PCA acelera el fit (~0.72 s → ~0.3 s), pero la SVD que requiere (~29 s) y la mayor RAM anulan la "
    "ganancia: el cuello de botella es la carga del CSV, no el ajuste."
)
p(
    "se conserva Ridge sin PCA como modelo final: mejores métricas, menos memoria y "
    "arquitectura más simple. PCA queda como parámetro reproducible (-use-pca).",
    bold_prefix="Decisión: ",
)

h("5. Arquitectura (resumen)", 2)
p(
    "El sistema en Go opera en dos bloques. (1) Pipeline concurrente de ingesta: una goroutine lee el "
    "CSV, un pool de workers parsea y valida, y otro pool codifica las categóricas; las etapas se "
    "comunican por canales con buffer y se sincronizan con WaitGroups, contadores atómicos y context "
    "con cancelación. (2) Entrenamiento paralelo: las filas se reparten entre fit-workers que acumulan "
    "parciales locales de XᵀX y Xᵀy (map-reduce sin escrituras compartidas), se reducen tras una "
    "barrera y se resuelve por Cholesky con respaldo SVD. Toda la configuración es parametrizable por "
    "CLI, lo que hace cada experimento reproducible con un solo comando."
)

out = "/Users/azariel/Documents/Github/PC3/metodologia_datos_y_modelo.docx"
doc.save(out)
print(f"OK: {out}")
